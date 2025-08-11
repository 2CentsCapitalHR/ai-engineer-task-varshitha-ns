"""Document type classification module for ADGM documents"""

import re
from typing import Dict, List, Tuple, Optional
from docx import Document
import logging
try:
    from .config import DOCUMENT_TYPE_MAPPINGS, PROCESS_CHECKLISTS
except ImportError:
    from config import DOCUMENT_TYPE_MAPPINGS, PROCESS_CHECKLISTS

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ADGMDocumentClassifier:
    """Classifies uploaded documents based on content analysis"""
    
    def __init__(self):
        self.document_mappings = DOCUMENT_TYPE_MAPPINGS
        self.process_checklists = PROCESS_CHECKLISTS
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text content from DOCX file"""
        try:
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            return " ".join(text_content).lower()
        
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""
    
    def classify_document_type(self, file_path: str, filename: str) -> str:
        """Classify document type based on content and filename"""

        # First, try to classify based on filename with scoring for better matches
        filename_lower = filename.lower()

        # Score each document type based on keyword matches (longer matches get higher scores)
        scores = {}
        for doc_type, keywords in self.document_mappings.items():
            score = 0
            best_match = ""
            for keyword in keywords:
                if keyword in filename_lower:
                    # Longer, more specific keywords get higher scores
                    keyword_score = len(keyword.replace(" ", ""))
                    if keyword_score > score:
                        score = keyword_score
                        best_match = keyword
            scores[doc_type] = (score, best_match)

        # Return the document type with the highest score
        if scores:
            best_doc_type = max(scores, key=lambda x: scores[x][0])
            if scores[best_doc_type][0] > 0:
                logger.info(f"Classified {filename} as {best_doc_type} based on filename keyword '{scores[best_doc_type][1]}'")
                return best_doc_type
        
        # If filename classification fails, analyze content
        text_content = self.extract_text_from_docx(file_path)
        
        if not text_content:
            return "unknown"
        
        # Score each document type based on keyword matches
        scores = {}
        for doc_type, keywords in self.document_mappings.items():
            score = 0
            for keyword in keywords:
                # Count occurrences of each keyword
                occurrences = len(re.findall(r'\b' + re.escape(keyword) + r'\b', text_content))
                score += occurrences
            scores[doc_type] = score
        
        # Return the document type with the highest score
        if scores and max(scores.values()) > 0:
            best_match = max(scores, key=scores.get)
            logger.info(f"Classified {filename} as {best_match} based on content analysis")
            return best_match
        
        # Additional specific patterns for better classification
        if "articles of association" in text_content or "company constitution" in text_content:
            return "articles_of_association"
        elif "memorandum of association" in text_content or "company objects" in text_content:
            return "memorandum_of_association"
        elif "board of directors" in text_content and "resolution" in text_content:
            return "board_resolution"
        elif "shareholders" in text_content and "resolution" in text_content:
            return "shareholder_resolution"
        elif "ultimate beneficial owner" in text_content or "ubo" in text_content:
            return "ubo_declaration"
        elif "employment" in text_content and ("contract" in text_content or "agreement" in text_content):
            return "employment_contract"
        
        logger.warning(f"Could not classify document: {filename}")
        return "unknown"
    
    def detect_process_type(self, document_types: List[str]) -> str:
        """Detect the business process based on uploaded document types"""
        
        # Count matches for each process type
        process_scores = {}
        
        for process_name, checklist in self.process_checklists.items():
            score = 0
            required_docs = checklist.get("required_documents", [])
            
            for doc_type in document_types:
                # Convert document type to readable format for matching
                readable_doc = self._convert_doc_type_to_readable(doc_type)
                
                for required_doc in required_docs:
                    if self._documents_match(readable_doc, required_doc):
                        score += 1
                        break
            
            process_scores[process_name] = score
        
        if process_scores and max(process_scores.values()) > 0:
            best_process = max(process_scores, key=process_scores.get)
            logger.info(f"Detected process type: {best_process}")
            return best_process
        
        return "Unknown Process"
    
    def _convert_doc_type_to_readable(self, doc_type: str) -> str:
        """Convert internal document type to readable format"""
        conversion_map = {
            "articles_of_association": "Articles of Association",
            "memorandum_of_association": "Memorandum of Association",
            "board_resolution": "Board Resolution",
            "shareholder_resolution": "Shareholder Resolution",
            "incorporation_form": "Incorporation Application Form",
            "ubo_declaration": "UBO Declaration Form",
            "register_members": "Register of Members and Directors",
            "employment_contract": "Employment Contract",
            "data_protection_policy": "Data Protection Policy",
            "annual_accounts": "Annual Accounts"
        }
        return conversion_map.get(doc_type, doc_type.replace("_", " ").title())
    
    def _documents_match(self, doc1: str, doc2: str) -> bool:
        """Check if two document names refer to the same document type"""
        doc1_lower = doc1.lower()
        doc2_lower = doc2.lower()
        
        # Direct match
        if doc1_lower == doc2_lower:
            return True
        
        # Fuzzy matching for common variations
        if "articles" in doc1_lower and "articles" in doc2_lower:
            return True
        if "memorandum" in doc1_lower and "memorandum" in doc2_lower:
            return True
        if "resolution" in doc1_lower and "resolution" in doc2_lower:
            return True
        if "ubo" in doc1_lower and ("ubo" in doc2_lower or "beneficial owner" in doc2_lower):
            return True
        if "register" in doc1_lower and "register" in doc2_lower:
            return True
        if "incorporation" in doc1_lower and ("incorporation" in doc2_lower or "application" in doc2_lower):
            return True
        
        return False
    
    def generate_checklist_report(self, document_types: List[str], process_type: str) -> Dict:
        """Generate a report on missing documents for the detected process"""
        
        if process_type not in self.process_checklists:
            return {
                "process": process_type,
                "status": "Unknown process type",
                "uploaded_documents": len(document_types),
                "missing_documents": [],
                "completeness_percentage": 0
            }
        
        checklist = self.process_checklists[process_type]
        required_docs = checklist.get("required_documents", [])
        
        # Convert uploaded document types to readable format
        uploaded_readable = [self._convert_doc_type_to_readable(doc_type) for doc_type in document_types]
        
        # Find missing documents
        missing_docs = []
        found_docs = []
        
        for required_doc in required_docs:
            found = False
            for uploaded_doc in uploaded_readable:
                if self._documents_match(uploaded_doc, required_doc):
                    found_docs.append(required_doc)
                    found = True
                    break
            
            if not found:
                missing_docs.append(required_doc)
        
        completeness = (len(found_docs) / len(required_docs)) * 100 if required_docs else 100
        
        return {
            "process": process_type,
            "status": "Complete" if not missing_docs else "Incomplete",
            "uploaded_documents": len(document_types),
            "required_documents": len(required_docs),
            "found_documents": found_docs,
            "missing_documents": missing_docs,
            "completeness_percentage": round(completeness, 1)
        }
