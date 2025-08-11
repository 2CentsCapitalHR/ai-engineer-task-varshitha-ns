"""Document processing module for DOCX inline commenting and review"""

import os
from typing import List, Dict, Any, Optional
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
import logging
from .document_classifier import ADGMDocumentClassifier
from .red_flag_detector import ADGMRedFlagDetector
from .rag_system import ADGMRAGSystem

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class ADGMDocumentProcessor:
    """Main document processor for ADGM compliance review"""
    
    def __init__(self):
        self.classifier = ADGMDocumentClassifier()
        self.red_flag_detector = ADGMRedFlagDetector()
        self.rag_system = ADGMRAGSystem()
    
    def process_documents(self, uploaded_files: List[Dict[str, str]], output_dir: str) -> Dict[str, Any]:
        """Process multiple uploaded documents and generate comprehensive review"""
        
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        results = {
            "summary": {},
            "documents": [],
            "process_analysis": {},
            "overall_compliance": {}
        }
        
        document_types = []
        all_issues = []
        
        # Process each document
        for file_info in uploaded_files:
            file_path = file_info["path"]
            filename = file_info["name"]
            
            logger.info(f"Processing document: {filename}")
            
            # Classify document type
            doc_type = self.classifier.classify_document_type(file_path, filename)
            document_types.append(doc_type)
            
            # Detect red flags and compliance issues
            analysis = self.red_flag_detector.analyze_document(file_path, doc_type)
            all_issues.extend(analysis.get("issues_found", []))

            # Get enhanced LLM analysis if available
            try:
                doc_text = self._extract_document_text(file_path)
                enhanced_analysis = self.rag_system.generate_enhanced_analysis(
                    doc_text, doc_type, analysis.get("issues_found", [])
                )
                analysis["enhanced_analysis"] = enhanced_analysis
            except Exception as e:
                logger.warning(f"Could not generate enhanced analysis: {str(e)}")
                analysis["enhanced_analysis"] = {"llm_available": False}
            
            # Create reviewed document with comments
            reviewed_doc_path = self._create_reviewed_document(
                file_path, filename, analysis, output_dir
            )
            
            # Document result
            doc_result = {
                "original_filename": filename,
                "document_type": doc_type,
                "reviewed_document_path": reviewed_doc_path,
                "analysis": analysis,
                "compliance_score": analysis.get("compliance_score", 0),
                "issues_count": len(analysis.get("issues_found", []))
            }
            
            results["documents"].append(doc_result)
        
        # Detect process type and check completeness
        process_type = self.classifier.detect_process_type(document_types)
        checklist_report = self.classifier.generate_checklist_report(document_types, process_type)
        
        # Generate overall summary
        results["summary"] = {
            "total_documents": len(uploaded_files),
            "process_type": process_type,
            "document_types_found": document_types,
            "total_issues": len(all_issues),
            "critical_issues": len([i for i in all_issues if i.get("severity") == "High"]),
            "warnings": len([i for i in all_issues if i.get("severity") == "Medium"])
        }
        
        results["process_analysis"] = checklist_report
        
        # Calculate overall compliance
        results["overall_compliance"] = self._calculate_overall_compliance(results["documents"])
        
        # Generate summary JSON
        self._save_summary_json(results, output_dir)
        
        return results
    
    def _create_reviewed_document(self, file_path: str, filename: str, analysis: Dict, output_dir: str) -> str:
        """Create a reviewed version of the document with inline comments"""
        try:
            # Load original document
            doc = Document(file_path)
            
            # Add review header
            self._add_review_header(doc, analysis)
            
            # Process issues and add comments
            issues = analysis.get("issues_found", [])
            
            for issue in issues:
                paragraph_index = issue.get("paragraph_index", 0)
                
                # Ensure paragraph index is valid
                if paragraph_index < len(doc.paragraphs):
                    paragraph = doc.paragraphs[paragraph_index]
                    
                    # Highlight the paragraph
                    self._highlight_paragraph(paragraph, issue["severity"])
                    
                    # Add comment as a new paragraph after the flagged one
                    comment_text = self._format_comment(issue)
                    self._add_comment_paragraph(doc, paragraph_index + 1, comment_text, issue["severity"])
            
            # Add compliance summary at the end
            self._add_compliance_summary(doc, analysis)
            
            # Save reviewed document
            output_filename = f"reviewed_{filename}"
            output_path = os.path.join(output_dir, output_filename)
            doc.save(output_path)
            
            logger.info(f"Created reviewed document: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating reviewed document for {filename}: {str(e)}")
            return ""
    
    def _add_review_header(self, doc: Document, analysis: Dict):
        """Add review header to the document"""
        # Insert header at the beginning
        header_para = doc.paragraphs[0]._element
        header_para.getparent().insert(0, header_para.getparent().makeelement(qn('w:p')))
        
        header = doc.paragraphs[0]
        header.text = "=== ADGM COMPLIANCE REVIEW ==="
        
        # Style the header
        header_run = header.runs[0]
        header_run.bold = True
        header_run.font.size = 240000  # 12pt in EMUs
        header_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        
        # Add summary info
        summary_para = doc.add_paragraph()
        summary_text = f"""
Document Type: {analysis.get('document_type', 'Unknown')}
Compliance Score: {analysis.get('compliance_score', 0)}/100
Total Issues Found: {analysis.get('total_flags', 0)}
Critical Issues: {analysis.get('critical_issues', 0)}
Warnings: {analysis.get('warnings', 0)}
Status: {analysis.get('compliance_status', 'Unknown')}

Review Date: {self._get_current_date()}
Generated by: ADGM Compliance Assistant

"""
        summary_para.text = summary_text
        summary_run = summary_para.runs[0]
        summary_run.font.size = 200000  # 10pt in EMUs
        
        # Add separator
        separator = doc.add_paragraph()
        separator.text = "=" * 80
        
        # Add space
        doc.add_paragraph()
    
    def _highlight_paragraph(self, paragraph, severity: str):
        """Highlight paragraph based on issue severity"""
        color_map = {
            "High": WD_COLOR_INDEX.RED,
            "Medium": WD_COLOR_INDEX.YELLOW,
            "Low": WD_COLOR_INDEX.BRIGHT_GREEN
        }
        
        highlight_color = color_map.get(severity, WD_COLOR_INDEX.YELLOW)
        
        for run in paragraph.runs:
            run.font.highlight_color = highlight_color
    
    def _format_comment(self, issue: Dict) -> str:
        """Format issue as a comment"""
        comment = f"""
🚨 COMPLIANCE ISSUE - {issue.get('severity', 'Unknown')} Severity

Issue: {issue.get('issue', 'Unknown issue')}
Found Pattern: {issue.get('pattern_found', 'N/A')}
Suggestion: {issue.get('suggestion', 'No suggestion available')}
ADGM Reference: {issue.get('adgm_reference', 'General ADGM Regulations')}

"""
        if issue.get('rag_guidance'):
            comment += f"Detailed Guidance:\n{issue['rag_guidance']}\n"
        
        comment += "-" * 50 + "\n"
        return comment
    
    def _add_comment_paragraph(self, doc: Document, insert_index: int, comment_text: str, severity: str):
        """Add a comment paragraph at specified index"""
        # Create new paragraph
        new_para = doc.add_paragraph()
        new_para.text = comment_text
        
        # Style the comment
        comment_run = new_para.runs[0]
        comment_run.font.size = 180000  # 9pt in EMUs
        comment_run.italic = True
        
        # Color based on severity
        if severity == "High":
            comment_run.font.color.rgb = RGBColor(220, 20, 60)  # Crimson
        elif severity == "Medium":
            comment_run.font.color.rgb = RGBColor(255, 140, 0)  # Dark orange
        else:
            comment_run.font.color.rgb = RGBColor(70, 130, 180)  # Steel blue
        
        # Move paragraph to correct position
        if insert_index < len(doc.paragraphs) - 1:
            new_para._element.getparent().insert(insert_index, new_para._element)
    
    def _add_compliance_summary(self, doc: Document, analysis: Dict):
        """Add compliance summary at the end of the document"""
        # Add separator
        doc.add_paragraph()
        separator = doc.add_paragraph()
        separator.text = "=" * 80
        
        # Add summary title
        summary_title = doc.add_paragraph()
        summary_title.text = "COMPLIANCE REVIEW SUMMARY"
        title_run = summary_title.runs[0]
        title_run.bold = True
        title_run.font.size = 220000  # 11pt
        title_run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        
        # Add detailed summary
        issues = analysis.get("issues_found", [])
        
        summary_text = f"""
Overall Compliance Score: {analysis.get('compliance_score', 0)}/100
Compliance Status: {analysis.get('compliance_status', 'Unknown')}

Issues Breakdown:
- Critical Issues (High): {analysis.get('critical_issues', 0)}
- Warnings (Medium): {analysis.get('warnings', 0)}
- Total Issues: {analysis.get('total_flags', 0)}

"""
        
        if issues:
            summary_text += "Issue Categories Found:\n"
            flag_types = list(set([issue.get('flag_type', 'Unknown') for issue in issues]))
            for flag_type in flag_types:
                count = len([i for i in issues if i.get('flag_type') == flag_type])
                summary_text += f"- {flag_type.replace('_', ' ').title()}: {count} issue(s)\n"
        
        summary_text += f"""
Next Steps:
1. Address all critical (High severity) issues immediately
2. Review and resolve warning (Medium severity) issues
3. Ensure all required documents are uploaded for the process
4. Re-submit for review after making corrections

Review completed on: {self._get_current_date()}
"""
        
        summary_para = doc.add_paragraph()
        summary_para.text = summary_text
        summary_run = summary_para.runs[0]
        summary_run.font.size = 200000  # 10pt
    
    def _calculate_overall_compliance(self, documents: List[Dict]) -> Dict:
        """Calculate overall compliance across all documents"""
        if not documents:
            return {"score": 0, "status": "No documents"}
        
        total_score = sum([doc.get("compliance_score", 0) for doc in documents])
        avg_score = total_score / len(documents)
        
        total_issues = sum([doc.get("issues_count", 0) for doc in documents])
        
        # Determine overall status
        if avg_score >= 90:
            status = "Compliant"
        elif avg_score >= 70:
            status = "Minor Issues"
        elif avg_score >= 50:
            status = "Major Issues"
        else:
            status = "Non-Compliant"
        
        return {
            "average_score": round(avg_score, 1),
            "status": status,
            "total_issues": total_issues,
            "documents_reviewed": len(documents)
        }
    
    def _save_summary_json(self, results: Dict, output_dir: str):
        """Save summary results as JSON"""
        import json
        
        # Create a simplified version for JSON output
        json_summary = {
            "process": results["process_analysis"].get("process", "Unknown"),
            "documents_uploaded": results["summary"]["total_documents"],
            "required_documents": results["process_analysis"].get("required_documents", 0),
            "missing_documents": results["process_analysis"].get("missing_documents", []),
            "completeness_percentage": results["process_analysis"].get("completeness_percentage", 0),
            "overall_compliance_score": results["overall_compliance"].get("average_score", 0),
            "overall_status": results["overall_compliance"].get("status", "Unknown"),
            "issues_found": []
        }
        
        # Add simplified issues
        for doc in results["documents"]:
            for issue in doc["analysis"].get("issues_found", []):
                simplified_issue = {
                    "document": doc["original_filename"],
                    "document_type": doc["document_type"],
                    "issue": issue.get("issue", ""),
                    "severity": issue.get("severity", ""),
                    "suggestion": issue.get("suggestion", ""),
                    "adgm_reference": issue.get("adgm_reference", "")
                }
                json_summary["issues_found"].append(simplified_issue)
        
        # Save JSON file
        json_path = os.path.join(output_dir, "compliance_summary.json")
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_summary, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved compliance summary to: {json_path}")
    
    def _extract_document_text(self, file_path: str) -> str:
        """Extract text content from document for LLM analysis"""
        try:
            doc = Document(file_path)
            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())

            return "\n".join(text_content)

        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return ""

    def _get_current_date(self) -> str:
        """Get current date formatted"""
        from datetime import datetime
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")
